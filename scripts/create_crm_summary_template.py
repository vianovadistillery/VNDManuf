"""Create CRM_Summary.docx template for customer CRM PDF export.

Run from project root: python scripts/create_crm_summary_template.py

docxtpl: {%tr for %} / {%tr endfor %} must live in the same table row XML part.
Do not wrap tables in {% if %} paragraphs — section data is filtered in Python.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "templates" / "CRM_Summary.docx"


def _add_loop_table(doc, headers: list[str], row_cells: list[str]) -> None:
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data_row = table.rows[1].cells
    for i, tpl in enumerate(row_cells):
        data_row[i].text = tpl
    doc.add_paragraph()


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = doc.add_heading("CRM Customer Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Customer: {{ contact.name }}")
    doc.add_paragraph("Code: {{ contact.code }}")
    doc.add_paragraph("Document: {{ document.doc_number }}   Date: {{ document.date }}")
    doc.add_paragraph(
        "Period: {{ document.period_start }} to {{ document.period_end }}"
    )
    doc.add_paragraph()

    doc.add_heading("Profile", level=1)
    doc.add_paragraph("Buying group: {{ profile.buying_group | default('—') }}")
    doc.add_paragraph(
        "Visit target (days): {{ profile.visit_target_days | default('—') }}"
    )
    doc.add_paragraph(
        "Preferred contact: {{ profile.preferred_contact | default('—') }}"
    )
    doc.add_paragraph("Primary rep: {{ profile.primary_rep | default('—') }}")
    doc.add_paragraph("Type: {{ profile.customer_type | default('—') }}")
    doc.add_paragraph("Email: {{ profile.email | default('—') }}")
    doc.add_paragraph("Phone: {{ profile.phone | default('—') }}")
    doc.add_paragraph("Suggestion: {{ profile.suggestion_message | default('') }}")
    doc.add_paragraph()

    doc.add_heading("Sales summary", level=1)
    doc.add_paragraph(
        "Orders: {{ document.order_count }}   "
        "Total ex GST: {{ document.tot_ex }}   "
        "Total inc GST: {{ document.tot }}   "
        "Units: {{ document.total_ordered }}"
    )
    doc.add_heading("Orders", level=2)
    _add_loop_table(
        doc,
        ["Date", "Ref", "PO", "Status", "Total ex", "Total inc"],
        [
            "{% for o in order_items %}{{ o.order_date }}",
            "{{ o.order_ref }}",
            "{{ o.po_number }}",
            "{{ o.status }}",
            "{{ o.total_ex_gst }}",
            "{{ o.total_inc_gst }}{% endfor %}",
        ],
    )
    doc.add_heading("SKU summary", level=2)
    _add_loop_table(
        doc,
        ["SKU", "Product", "Qty", "Unit ex", "Line inc"],
        [
            "{% for row in line_items %}{{ row.sku }}",
            "{{ row.product_name or row.description }}",
            "{{ row.quantity }}",
            "{{ row.ex }}",
            "{{ row.tot_inc }}{% endfor %}",
        ],
    )
    doc.add_paragraph()

    doc.add_heading("Timeline", level=1)
    _add_loop_table(
        doc,
        ["When", "Type", "Title", "Source"],
        [
            "{% for item in timeline_items %}{{ item.when }}",
            "{{ item.type }}",
            "{{ item.title }}",
            "{{ item.source }}{% endfor %}",
        ],
    )
    doc.add_paragraph()

    doc.add_heading("People", level=1)
    _add_loop_table(
        doc,
        ["Name", "Role", "Phone", "Email", "Primary", "Notes"],
        [
            "{% for s in staff_items %}{{ s.name }}",
            "{{ s.role }}",
            "{{ s.phone }}",
            "{{ s.email }}",
            "{{ s.is_primary }}",
            "{{ s.notes }}{% endfor %}",
        ],
    )
    doc.add_paragraph()

    doc.add_heading("Scheduled activities", level=1)
    _add_loop_table(
        doc,
        ["When", "Type", "Title", "Status", "Description"],
        [
            "{% for s in scheduled_items %}{{ s.when }}",
            "{{ s.type }}",
            "{{ s.title }}",
            "{{ s.status }}",
            "{{ s.description }}{% endfor %}",
        ],
    )

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
