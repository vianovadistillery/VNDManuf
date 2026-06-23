"""Create Customer_Purchase_Report.docx template (Vianova style, based on Delivery_Docket layout).

Run from project root: python scripts/create_customer_purchase_report_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates"
OUTPUT_PATH = TEMPLATES_DIR / "Customer_Purchase_Report.docx"


def main():
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CUSTOMER ")
    run.bold = True
    run.font.size = Pt(18)
    run2 = title.add_run("PURCHASE SUMMARY")
    run2.bold = True
    run2.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph("Customer: {{ contact.name }}")
    doc.add_paragraph("Code: {{ contact.code }}")
    doc.add_paragraph("ABN: {{ contact.abn }}")
    doc.add_paragraph(
        "Period: {{ document.period_start }} to {{ document.period_end }}"
    )
    doc.add_paragraph("Orders in period: {{ document.order_count }}")
    doc.add_paragraph(
        "Report date: {{ document.date }}    Reference: {{ document.doc_number }}"
    )
    doc.add_paragraph()

    doc.add_paragraph("Orders")
    orders_table = doc.add_table(rows=3, cols=6)
    orders_table.style = "Table Grid"
    oh = orders_table.rows[0].cells
    oh[0].text = "Date"
    oh[1].text = "Order Ref"
    oh[2].text = "PO"
    oh[3].text = "Status"
    oh[4].text = "Total Ex"
    oh[5].text = "Total Inc"
    orow = orders_table.rows[1].cells
    orow[0].text = "{% for order in order_items %}{{ order.order_date }}"
    orow[1].text = "{{ order.order_ref }}"
    orow[2].text = "{{ order.po_number }}"
    orow[3].text = "{{ order.status }}"
    orow[4].text = "{{ order.total_ex_gst }}"
    orow[5].text = "{{ order.total_inc_gst }}"
    oend = orders_table.rows[2].cells
    oend[0].text = "{% endfor %}"

    doc.add_paragraph()
    doc.add_paragraph("Products in filtered orders")

    products_table = doc.add_table(rows=3, cols=6)
    products_table.style = "Table Grid"
    header = products_table.rows[0].cells
    header[0].text = "Code"
    header[1].text = "Description"
    header[2].text = ""
    header[3].text = "Unit"
    header[4].text = "Qty"
    header[5].text = "Total Ex"

    row1 = products_table.rows[1].cells
    row1[0].text = "{% for item in line_items %}{{ item.sku }}"
    row1[1].text = "{{ item.description }}"
    row1[2].text = ""
    row1[3].text = "{{ item.uom }}"
    row1[4].text = "{{ item.oqty }}"
    row1[5].text = "{{ item.tot_ex }}"

    row2 = products_table.rows[2].cells
    row2[0].text = "{% endfor %}"
    row2[1].text = ""
    row2[2].text = ""
    row2[3].text = "Total"
    row2[4].text = "{{ document.total_ordered }}"
    row2[5].text = "{{ document.tot_ex }}"

    doc.add_paragraph()
    doc.add_paragraph(
        "Total ex GST: {{ document.tot_ex }}    GST: {{ document.tot_gst }}    Total inc GST: {{ document.tot }}"
    )
    doc.add_paragraph("Notes: {{ document.notes }}")

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
