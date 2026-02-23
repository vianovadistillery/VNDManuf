"""Create a minimal delivery_docket.docx template with docxtpl placeholders.
Run from project root: python scripts/create_delivery_docket_template.py
"""

from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates"
OUTPUT_PATH = TEMPLATES_DIR / "delivery_docket.docx"


def main():
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Delivery Docket", 0)
    doc.add_paragraph()
    # Single-value placeholders (docxtpl/Jinja2)
    doc.add_paragraph("Customer: {{ contact.name }}")
    doc.add_paragraph("Code: {{ contact.code }}")
    doc.add_paragraph("Document: {{ document.doc_number }}  Date: {{ document.date }}")
    doc.add_paragraph("Delivery date: {{ document.delivery_date }}")
    doc.add_paragraph()
    doc.add_paragraph("Notes: {{ document.notes }}")
    doc.add_paragraph()
    # Table: header row + one data row with loop
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Description"
    header[1].text = "SKU"
    header[2].text = "Qty"
    header[3].text = "Unit Price"
    header[4].text = "Line Total"
    row1 = table.rows[1].cells
    # docxtpl: for/endfor around the row so one row repeats per item
    row1[0].text = "{% for item in line_items %}{{ item.description }}"
    row1[1].text = "{{ item.sku }}"
    row1[2].text = "{{ item.quantity }}"
    row1[3].text = "{{ item.unit_price }}"
    row1[4].text = "{{ item.line_total }}{% endfor %}"
    doc.add_paragraph()
    doc.add_paragraph("Subtotal: {{ document.subtotal }}   Total: {{ document.total }}")
    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
